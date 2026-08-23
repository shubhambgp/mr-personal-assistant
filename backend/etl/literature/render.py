"""Renders the synthetic corpus: Markdown sources, then PDF and DOCX from them.

Markdown is committed as the source of record because it is diffable in a pull
request; the PDF and DOCX are committed too, so CI and any reviewer can run the
real parser against real binary files rather than a convenient text fixture.

Section numbering follows the SmPC convention (4.1 Indications, 4.2 Posology,
4.5 Interactions …) on purpose: the numbers become the citation a rep reads, and
"Cardevia §4.5" is a reference somebody can check. Dosing is rendered as a real
table, which is exactly where pypdf is weakest — see etl/docs_parse.py.
"""

from __future__ import annotations

import textwrap
from pathlib import Path

FICTION_NOTICE = (
    "FICTIONAL DEMONSTRATION DOCUMENT — Qorvexa Healthcare is an invented company. "
    "This product, its molecule and every clinical statement here are fabricated for "
    "software testing. Not medical information. Do not use for any clinical purpose."
)


def monograph_markdown(brand: str, d: dict) -> str:
    """One monograph as Markdown, with the sections that brand actually has."""
    out: list[str] = []
    add = out.append

    add(f"# {brand} ({d['molecule']}) {' / '.join(d['strengths'])} {d['form']}")
    add("")
    add(f"> {FICTION_NOTICE}")
    add("")
    add(f"**Qorvexa Healthcare · Summary of Product Characteristics · {d['area']}**")
    add("")
    add("## 1. Name of the medicinal product")
    add("")
    add(f"{brand} {' / '.join(d['strengths'])} {d['form']}.")
    add("")
    add("## 2. Qualitative and quantitative composition")
    add("")
    add(f"Each {d['form']} contains {d['molecule']} as the active substance. "
        f"{brand} is a {d['class']}.")
    add("")
    add("## 4.1 Therapeutic indications")
    add("")
    for line in d["indications"]:
        add(f"- {line}")
    add("")
    add("## 4.2 Posology and method of administration")
    add("")
    add("| Population | Dose | Notes |")
    add("|---|---|---|")
    for population, dose, note in d["dosing"]:
        add(f"| {population} | {dose} | {note} |")
    add("")
    add("### 4.2.1 Renal impairment")
    add("")
    add(d["renal"])
    add("")
    add("### 4.2.2 Hepatic impairment")
    add("")
    add(d["hepatic"])
    add("")
    if "paediatric" in d:
        add("### 4.2.3 Paediatric population")
        add("")
        add(d["paediatric"])
        add("")
    add("## 4.3 Contraindications")
    add("")
    for line in d["contraindications"]:
        add(f"- {line}")
    add("")
    add("## 4.5 Interaction with other medicinal products")
    add("")
    add("| Interacting substance | Effect and action required |")
    add("|---|---|")
    for substance, effect in d["interactions"]:
        add(f"| {substance} | {effect} |")
    add("")
    if "pregnancy" in d:
        add("## 4.6 Fertility, pregnancy and lactation")
        add("")
        add(d["pregnancy"])
        add("")
    add("## 4.8 Undesirable effects")
    add("")
    for frequency, effects in d["adverse"]:
        add(f"- **{frequency}:** {effects}")
    add("")
    add("## 6.4 Special precautions for storage")
    add("")
    add(d["storage"])
    add("")
    return "\n".join(out)


# --------------------------------------------------------------------- PDF ---

def markdown_to_pdf(markdown: str, out_path: Path, title: str) -> int:
    """Renders Markdown to PDF with real headings and real tables.

    Returns the page count. Tables are rendered as tables rather than flattened
    to text on purpose: a monograph's dosing table is the single most
    safety-relevant thing in it, and a corpus that avoided tables would let a
    parser weakness go unnoticed.
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    base = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=base["BodyText"], fontSize=9.5, leading=13,
                          alignment=TA_LEFT, spaceAfter=4)
    h1 = ParagraphStyle("h1", parent=base["Heading1"], fontSize=15, leading=19, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=base["Heading2"], fontSize=11.5, leading=15,
                        spaceBefore=10, spaceAfter=4)
    h3 = ParagraphStyle("h3", parent=base["Heading3"], fontSize=10.5, leading=14,
                        spaceBefore=8, spaceAfter=3)
    notice = ParagraphStyle("notice", parent=body, fontSize=8, textColor=colors.HexColor("#9a3412"),
                            backColor=colors.HexColor("#fff7ed"), borderPadding=5, spaceAfter=8)
    cell = ParagraphStyle("cell", parent=body, fontSize=8.5, leading=11, spaceAfter=0)

    story: list = []
    pending_rows: list[list[str]] = []

    def flush_table() -> None:
        if not pending_rows:
            return
        header, *rows = pending_rows
        data = [[Paragraph(f"<b>{c}</b>", cell) for c in header]]
        data += [[Paragraph(c, cell) for c in r] for r in rows]
        widths = [165 * mm / len(header)] * len(header)
        table = Table(data, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(table)
        story.append(Spacer(1, 6))
        pending_rows.clear()

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # the markdown separator row
            pending_rows.append(cells)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("> "):
            story.append(Paragraph(line[2:], notice))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], h3))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], h2))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], h1))
        elif line.startswith("- "):
            story.append(Paragraph(f"• {_inline(line[2:])}", body))
        else:
            story.append(Paragraph(_inline(line), body))
    flush_table()

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4, title=title, author="Qorvexa Healthcare (fictional)",
        leftMargin=22 * mm, rightMargin=22 * mm, topMargin=20 * mm, bottomMargin=20 * mm,
    )

    def footer(canvas, document) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 6.5)
        canvas.setFillColor(colors.HexColor("#94a3b8"))
        canvas.drawString(22 * mm, 12 * mm, FICTION_NOTICE[:110])
        canvas.drawRightString(A4[0] - 22 * mm, 12 * mm, f"Page {document.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    from pypdf import PdfReader

    return len(PdfReader(str(out_path)).pages)


def _inline(text: str) -> str:
    """Markdown bold/italic to reportlab's mini-HTML."""
    import re

    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", text)


# -------------------------------------------------------------------- DOCX ---

def markdown_to_docx(markdown: str, out_path: Path, title: str) -> None:
    """Detailing aids arrive as Word documents in real life, so some do here."""

    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.core_properties.title = title
    document.core_properties.comments = FICTION_NOTICE

    pending: list[list[str]] = []

    def flush() -> None:
        if not pending:
            return
        header, *rows = pending
        table = document.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"
        for i, text in enumerate(header):
            run = table.rows[0].cells[i].paragraphs[0].add_run(text)
            run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, text in enumerate(row):
                cells[i].text = _plain(text)
        document.add_paragraph()
        pending.clear()

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue
            pending.append(cells)
            continue
        flush()
        if not line:
            continue
        if line.startswith("> "):
            para = document.add_paragraph()
            run = para.add_run(_plain(line[2:]))
            run.italic = True
            run.font.size = Pt(8)
        elif line.startswith("### "):
            document.add_heading(_plain(line[4:]), level=3)
        elif line.startswith("## "):
            document.add_heading(_plain(line[3:]), level=2)
        elif line.startswith("# "):
            document.add_heading(_plain(line[2:]), level=1)
        elif line.startswith("- "):
            document.add_paragraph(_plain(line[2:]), style="List Bullet")
        else:
            document.add_paragraph(_plain(line))
    flush()
    document.save(str(out_path))


def _plain(text: str) -> str:
    import re

    return re.sub(r"\*{1,2}", "", text)


def wrap(text: str, width: int = 88) -> str:
    return "\n".join(textwrap.fill(p, width) for p in text.split("\n"))
